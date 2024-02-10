
from tqdm import tqdm
for i in tqdm(range(1), desc='Loading...', ncols=75):
    import os
    import subprocess
    import pandas
    from textacy import extract
    import spacy
    import pyinputplus
    from pdfminer.high_level import extract_text
    from docx2python import docx2python
    from pathlib import Path
    from docx import Document
    from docx.shared import Inches, Pt
    from accessory import text_cleanup, pdf_cleanup, highlight, colourise

global text, case_sensitive, search_phrase, search_phrase_list, master_list, sentences, party, target_File_path, \
    result_filename


def kwic(text, party):
    result = extract.kwic.keyword_in_context(text, keyword=party, window_width=80)
    df = pandas.DataFrame(result, columns=['Left', 'Key', 'Right'])
    pandas.set_option('display.max_rows', None)
    pandas.set_option('display.max_colwidth', None)
    pandas.set_option('display.max_columns', 3)
    result_file = rf'{target_File_path}\{result_filename}_{party}_search_result.xlsx'
    df.to_excel(result_file, index=False)
    subprocess.Popen([r'C:\Program Files\Microsoft Office\root\Office16\EXCEL.EXE', result_file])


def is_match(sent):
    for i in search_phrase_list:
        if i in sent:
            return sent


while True:
    print("")
    file = input('Enter the file path and name - only .docx, .pdf or .txt files: ')
    print("")
    target_File = file.strip('\"')
    filename, file_extension = os.path.splitext(target_File)
    p = Path(target_File)
    target_File_path = p.parent
    result_filename = p.name
    extensions = ['.docx', '.pdf', '.txt']

    if file_extension not in extensions:
        print('Invalid file type.')
        continue

    if file_extension in extensions:
        for i in tqdm(range(1), desc='File processing', ncols=75):
            if file_extension == '.docx':
                docx_content = docx2python(target_File)
                text = docx_content.text
                text = text_cleanup(text)
            if file_extension == '.pdf':
                text = extract_text(target_File)
                text = text_cleanup(text)
                text = pdf_cleanup(text)
            if file_extension == '.txt':
                target = open(target_File, 'r', encoding='utf8')
                text = target.read()
                text = text_cleanup(text)
        break

print('\n', 'Is this search case sensitive?')
case_sensitive = pyinputplus.inputMenu(['Yes', 'No'], numbered=True)
if case_sensitive == 'No':
    text = text.lower()
print("")

for i in tqdm(range(1), desc='Sentence tokenization', ncols=75):
    nlp = spacy.load("en_core_web_lg")
    doc = nlp(text)
    sentences: list[str] = [sentence.text for sentence in doc.sents]
print("")

while True:
    master_List = []
    search_phrase_list: list[str] = []

    party = pyinputplus.inputStr(r'''Enter the term for the party to be searched (entry is case sensitive if that
     option selected):''')
    if case_sensitive == 'No':
        party = party.lower()
    print("")

    for i in tqdm(range(1), desc='KWIC processing', ncols=75):
        kwic(text, party)
    print("")

    while True:
        secTerm = pyinputplus.inputStr(r'''Enter a predicate search term or phrase (include "'s" or "'" for possessive
         case of the party being searched):''')

        if not secTerm.startswith("'"):  # add a space where the second term is not an apostrophe
            secTerm = " ".join(["", secTerm])

        search_phrase = party + secTerm
        if case_sensitive == 'No':
            search_phrase = search_phrase.lower()

        search_phrase_list.append(search_phrase)
        list_Status = ' | '.join(search_phrase_list)
        print('')
        print(highlight("blue", (colourise("gray", f'Search Phrases: {list_Status}'))))  # tracks Search Phrases
        print('')

        response = pyinputplus.inputMenu(['Enter another search term', 'Finished for this party'], numbered=True)

        if response == 'Finished for this party':
            master_List = [sent for sent in sentences if is_match(sent)]
            break

    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    sections = document.sections
    section = sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    document.add_paragraph(f'Phrases searched: {list_Status}')
    for j in master_List:
        document.add_paragraph(j)
    document.save(rf'{target_File_path}\{result_filename}_{party}_search_result.docx')
    subprocess.Popen([r'C:\Program Files\Microsoft Office\root\Office16\WINWORD.EXE',
                      rf'{target_File_path}\{result_filename}_{party}_search_result.docx'])
    response = pyinputplus.inputMenu(['Search another party', 'Finished'], numbered=True)
    if response == 'Search another party':
        continue
    if response == 'Finished':
        print('Run Finished')
    break
