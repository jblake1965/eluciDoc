

from rich.progress import track
import time
for i in track(range(2), description="[bold magenta italic]Loading...[/]"):
    time.sleep(1)
    import re
    import os
    import subprocess
    import pandas
    from textacy import extract
    import spacy
    from pdfminer.high_level import extract_text
    from docx2python import docx2python
    from pathlib import Path
    from docx import Document
    from docx.shared import Inches, Pt
    from accessory import text_cleanup, pdf_cleanup
    from rich.console import Console
    console = Console()


global text, case_sensitive, search_phrase, search_phrase_list, master_list, sentences, party, target_File_path, \
    result_filename


def kwic(text, party):
    global result_file
    result = extract.kwic.keyword_in_context(text, keyword=party, window_width=80)
    df = pandas.DataFrame(result, columns=['Left', 'Key', 'Right'])
    pandas.set_option('display.max_rows', None)
    pandas.set_option('display.max_colwidth', None)
    pandas.set_option('display.max_columns', 3)
    try:
        result_file = rf'{target_File_path}\{result_filename}_{party}_search_result.xlsx'
        df.to_excel(result_file, index=False)
        subprocess.Popen([r'C:\Program Files\Microsoft Office\root\Office16\EXCEL.EXE', result_file])
    except PermissionError:
        print('\n')
        console.print(rf'''[bold green italic]***NOTE: AN INSTANCE OF [/][bold red italic]"{target_File_path}\{result_filename}_{party}_search_result.xlsx" [/][bold green italic]IS ALREADY OPEN***[/]''')
        # fix so it prompts to close open instance of file


def is_match(sent):
    for i in search_phrase_list:
        if i in sent:
            return sent


while True:
    console.print("")
    file = console.input('[bold green italic]ENTER THE FILE PATH AND NAME - ONLY .DOCX, .PDF or .TXT FILES:[/]')
    console.print("")
    target_File = file.strip('\"')
    filename, file_extension = os.path.splitext(target_File)
    p = Path(target_File)
    target_File_path = p.parent
    result_filename = p.name
    extensions = ['.docx', '.pdf', '.txt']

    if file_extension not in extensions:
        console.print('[bold red italic]INVALID FILE TYPE - PLEASE RE-ENTER.[/]')
        continue

    if file_extension in extensions:
        for i in track(range(2), description="[bold magenta italic]Processing text...[/]"):
            time.sleep(1)
            if file_extension == '.docx':
                docx_content = docx2python(target_File)
                text = docx_content.text
                text = text_cleanup(text)
            if file_extension == '.pdf':
                text = extract_text(target_File)
                text = text_cleanup(text)
                text = pdf_cleanup(text)
            if file_extension == '.txt':
                target = open(target_File, 'r', encoding='utf8')
                text = target.read()
                text = text_cleanup(text)
        break

while True:
    print("")
    case_sensitive = console.input('[bold green italic]IS THIS SEARCH CASE SENSITIVE?: YES or NO [/]')
    print("")
    case_sensitive = case_sensitive.lower()
    if case_sensitive not in ['yes', 'no']:
        console.print('\n', r'[bold red italic]PLEASE ENTER "YES" or "NO"[/]', '\n')
        continue
    if case_sensitive == 'no':
        text = text.lower()
        break
    if case_sensitive == 'yes':
        break

for i in track(range(2), description="[bold magenta italic]Tokenizing sentences...[/]"):
    time.sleep(1)
    nlp = spacy.load("en_core_web_lg")
    doc = nlp(text)
    sentences: list[str] = [sentence.text for sentence in doc.sents]
console.print("")

while True:
    master_List = []
    search_phrase_list: list[str] = []

    party = console.input(r'''[bold green italic]ENTER THE TERM FOR THE PARTY BEING SEARCHED (ENTRY IS CASE SENSITIVE IF
     THAT OPTION IS SELECTED):[/]''')
    if party == "":
        console.print('\n', r'[bold red italic]***PARTY NAME MUST BE ENTERED***[/]', '\n')
        continue

    if case_sensitive == 'No':
        party = party.lower()
    console.print("")

    kwic(text, party)
    console.print("")

    while True:
        secTerm = console.input(r'''[bold green italic]ENTER A PREDICATE SEARCH TERM OR PHRASE (INCLUDE "'S" OR "'" 
        FOR THE POSSESSIVE CASE OF THE PARTY BEING SEARCHED):[/]''')

        if secTerm == "":
            console.print('\n', r'[bold red italic]***A PREDICATE SEARCH TERM OR PHRASE MUST BE ENTERED***[/]', '\n')
            continue

        valid = re.compile(r"^[',;:)].*")
        if not valid.match(rf"{secTerm}"):
            secTerm = " ".join(["", secTerm])

        search_phrase = party + secTerm
        if case_sensitive == 'No':
            search_phrase = search_phrase.lower()

        search_phrase_list.append(search_phrase)
        list_Status = ' | '.join(search_phrase_list)
        console.print('')
        console.print(f'[bold white on blue]Search Phrases: [italic]{list_Status}[/]')  # tracks Search Phrases
        console.print('')

        while True:
            add_term_response = console.input('[bold green italic]ENTER ANOTHER SEARCH TERM?: YES or NO [/]')
            print("")
            add_term_response = add_term_response.lower()
            if add_term_response not in ['yes', 'no']:
                console.print('\n', r'[bold red italic]PLEASE ENTER "YES" or "NO"[/]', '\n')
                continue
            else:
                break

        if add_term_response == 'no':
            master_List = [sent for sent in sentences if is_match(sent)]
            break

    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    sections = document.sections
    section = sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    document.add_paragraph(f'Phrases searched: {list_Status}')
    for j in master_List:
        document.add_paragraph(j)
    try:
        document.save(rf'{target_File_path}\{result_filename}_{party}_search_result.docx')
    except PermissionError:
        console.print('\n', rf'''[bold red italic]***AN INSTANCE OF "{target_File_path}\{result_filename}_{party}_search_result.docx" IS 
        ALREADY OPEN. CLOSE FILE AND PRESS <ENTER>***''')
        wait = input("")
    subprocess.Popen([r'C:\Program Files\Microsoft Office\root\Office16\WINWORD.EXE',
                      rf'{target_File_path}\{result_filename}_{party}_search_result.docx'])

    while True:
        another_search_response = console.input('[bold green italic]RUN A SEARCH FOR ANOTHER PARTY?: YES or NO [/]')
        print("")
        another_search_response = another_search_response.lower()
        if another_search_response not in ['yes', 'no']:
            console.print('\n', r'[bold red italic]PLEASE ENTER "YES" or "NO"[/]', '\n')
            continue
        else:
            break

    if another_search_response == 'yes':
        continue

    if another_search_response == 'no':
        console.print('[bold magenta italic]RUN FINISHED[/]')
        break
    break
